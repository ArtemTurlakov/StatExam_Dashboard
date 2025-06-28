from fastapi import APIRouter
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

from ..router import CACHE

router = APIRouter(prefix="/report", tags=["Отчет"])

def doc_table(doc, data, type, subtype, label):
    cols = len(data)
    rows = len(data[0][type])
    table = doc.add_table(rows=rows+2, cols=cols*2+1, style='Table Grid')

    table.cell(1, 0).text = label

    for i in range(rows):
        table.cell(i+2, 0).text = dict(data[0][type][i])[subtype]
        table.cell(i+2, 0).width = Inches(3)
    for i in range(rows+2):
        table.cell(i, 2).width = Inches(2)
        table.cell(i, 4).width = Inches(2)
    for i in range(cols):
        table.cell(0, i*2+1).text = f"{data[i]['year']}"

        table.cell(1, i*2+1).text = "чел."
        table.cell(1, i*2+2).text = f"% от общего числа уч."
        for j in range(rows):
            table.cell(j+2, i*2+1).text = f"{dict(data[0][type][j])['count']}"
            table.cell(j+2, i*2+2).text = f"% .2f" % (dict(data[0][type][j])['count'] / data[i]['count_on_exam'])
    return table
    



@router.get("/{grade, subject, name, filename}",)
async def get_report(grade, subject, name, filename):
    data = CACHE.get(f'{grade}-{subject}', 0)
    cols = len(data)
    doc = Document("app/Dashboard/Report/Шаблон_отчета.docx")
    e = 'ОГЭ' if grade == "9" else 'ЕГЭ'
    doc.add_heading(f"Методический анализ результатов {e} \n по предмету {name}", 1)
    doc.add_heading(f"РАЗДЕЛ 1. ХАРАКТЕРИСТИКА УЧАСТНИКОВ {e} \n ПО УЧЕБНОМУ ПРЕДМЕТУ", 1)

    doc.add_paragraph("1.1.	Количество  участников экзамена по учебному предмету")
    table = doc.add_table(rows=3, cols=cols*2, style="Table Grid")
    for i in range(cols):
        table.cell(0, i*2).text = f"{data[i]['year']}"
        table.cell(1, i*2).text = "чел."
        table.cell(1, i*2+1).text = f"% от общего числа участников"
        table.cell(2, i*2).text = f"{data[i]['count_on_exam']}"
        table.cell(2, i*2+1).text = f"% .2f" % (data[i]['count_on_exam'] / data[i]['students_at_all'])

    doc.add_paragraph("1.2.	Процентное соотношение юношей и девушек, участвующих в экзамене")
    doc_table(doc, data, 'sextable', 'sex', "Пол")

    doc.add_paragraph("1.3.	Количество участников экзамена в регионе по категориям ")
    doc_table(doc, data, 'categories', 'description', "Категория участника")

    doc.add_paragraph("1.4.	Количество участников экзамена в регионе по типам  ОО ")
    doc_table(doc, data, 'OO', 'name', "Категория участника")

    doc.add_paragraph("1.5.	Количество участников ЕГЭ по учебному предмету по АТЕ региона")
    doc_table(doc, data, 'areas', 'name', "Наименование АТЕ")

    fig = plt.figure(figsize=(14, 5))
    fig.tight_layout()
    t = 'score' if grade == "9" else 'final_points'
    plt.bar([str(dict(x)[t]) for x in data[-1]['marks'][0]['finals_diag']],
             [dict(x)['count'] for x in data[-1]['marks'][0]['finals_diag']],
             width=0.8, )
    plt.xlabel('Баллы')
    plt.ylabel('Количество участников')
    
    plt.savefig('app/Dashboard/Report/diag.jpg', bbox_inches='tight')

    doc.add_paragraph(f"РАЗДЕЛ 2.  ОСНОВНЫЕ РЕЗУЛЬТАТЫ ЕГЭ ПО ПРЕДМЕТУ")
    doc.add_paragraph(f"2.1.	Диаграмма распределения тестовых баллов участников экзамена по предмету в {data[-1]['year']} г.")
    doc.add_picture('app/Dashboard/Report/diag.jpg', Inches(6))

    doc.save(f'app/Dashboard/Report/report.docx')

    return FileResponse(path=f'app/Dashboard/Report/report.docx',
                        filename=f'{filename}.docx',)
